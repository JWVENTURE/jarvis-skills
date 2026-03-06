"""
Professional Document Generator for JW Ventures
Use this script to generate invoices, work orders, and quotes.

Design Standards:
- Professional, clean, premium
- Olive Green (#8B9A6A) for JW Ventures branding (NOT PickleballBN colors!)
- Consistent footer on all pages
- Generous white space
- Side-by-side signatures for contracts
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# JW Ventures Company Info
COMPANY = {
    "name": "JW Ventures",
    "registration": "P30000657",
    "address": "No.30 Simpang 94, Kampong Pancha Delima, Jalan Muara",
    "city": "Bandar Seri Begawan",
    "email": "H.jabbar@jwventures.group",
    "website": "www.jwventures.group",
    "bank": "Bank of China (Hong Kong) Limited, Brunei Branch",
    "account": "052-120-2-004785-8",
    "swift": "BKCHBNBB",
    "owner": "Haji Awang Jabbar bin Haji Awang Tengah",
}

# JW Ventures Brand Colors (NOT PickleballBN colors!)
COLORS = {
    "olive": RGBColor(139, 154, 106),     # #8B9A6A - JW Ventures primary
    "charcoal": RGBColor(51, 51, 51),     # #333333 - Secondary
    "white": RGBColor(255, 255, 255),
    "gray_light": RGBColor(240, 240, 240), # #F0F0F0 - Alternating rows
}

# FOOTER text uses JW Ventures info
FOOTER_TEXT = f"JW Ventures | {COMPANY['registration']} | Bandar Seri Begawan | {COMPANY['email']}"


def set_cell_background(cell, color):
    """Set cell background color using hex or RGBColor"""
    shading_elm = OxmlElement('w:shd')
    if isinstance(color, str):
        shading_elm.set(qn('w:fill'), color)
    else:
        # RGBColor to hex
        r = color.rgbTriplet[0]
        g = color.rgbTriplet[1]
        b = color.rgbTriplet[2]
        hex_color = f"{r:02x}{g:02x}{b:02x}"
        shading_elm.set(qn('w:fill'), hex_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_colored_header_row(table, text, cols=1, background="8B9A6A", text_color=RGBColor(255, 255, 255)):
    """Add a header row with colored background"""
    if cols == 1:
        row = table.add_row().cells
        row[0].text = text
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.color.rgb = text_color
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(row[0], background)
    else:
        row = table.add_row().cells
        for i in range(cols):
            row[i].text = text
            row[i].paragraphs[0].runs[0].bold = True
            row[i].paragraphs[0].runs[0].font.color.rgb = text_color
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_background(row[i], background)


def add_footer(doc):
    """Add JW Ventures footer to all sections"""
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = FOOTER_TEXT
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)


def add_info_table(doc, info_dict, cols=4):
    """Add a formatted info table with key-value pairs"""
    table = doc.add_table(rows=len(info_dict), cols=cols)

    items = list(info_dict.items())
    for i, (key, value) in enumerate(items):
        row = i // cols
        col = (i % cols) * 2
        if col + 1 < cols:
            cell = table.rows[row].cells[col]
            cell.text = key
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = COLORS["olive"]
            cell.paragraphs[0].runs[0].font.size = Pt(10)

            cell = table.rows[row].cells[col + 1]
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    return table


def create_invoice(invoice_number, client_name, client_contact, items, due_date="Upon Receipt"):
    """Create a professional JW Ventures invoice"""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_heading('INVOICE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info Table
    doc.add_paragraph()
    info = {
        'Invoice No:': invoice_number,
        'Date:': 'March 7, 2026',
        'Client:': client_name,
        'Due Date:': due_date,
    }
    if client_contact:
        info['Contact:'] = client_contact

    table = doc.add_table(rows=2, cols=4)
    items = list(info.items())
    for i in range(4):
        if i < len(items):
            key, value = items[i]
            table.rows[0].cells[i].text = key
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = COLORS["olive"]
            table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

            table.rows[1].cells[i].text = str(value)
            table.rows[1].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Items Table with Styled Header
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = ['Description', 'Quantity', 'Unit Price', 'Amount (BND)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, '8B9A6A')
        cell.paragraphs[0].runs[0].font.color.rgb = COLORS["white"]

    # Data rows with alternating colors
    total = 0
    for i, item in enumerate(items):
        row_cells = table.add_row().cells
        row_cells[0].text = item["description"]
        row_cells[1].text = str(item.get("qty", 1))
        row_cells[2].text = item["unit_price"]
        row_cells[3].text = item["amount"]
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Alternating row color
        if i % 2 == 0:
            for cell in row_cells:
                set_cell_background(cell, 'F0F0F0')

        # Calculate total
        amt_str = item["amount"].replace("BND", "").strip()
        try:
            total += float(amt_str)
        except:
            pass

    # Total Row
    total_cells = table.add_row().cells
    total_cells[0].text = ''
    total_cells[1].text = ''
    total_cells[2].text = 'TOTAL'
    total_cells[2].paragraphs[0].runs[0].bold = True
    total_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_cells[3].text = f'BND {total:.0f}'
    total_cells[3].paragraphs[0].runs[0].bold = True
    total_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_background(total_cells[3], '8B9A6A')
    total_cells[3].paragraphs[0].runs[0].font.color.rgb = COLORS["white"]

    doc.add_paragraph()

    # Payment Details
    doc.add_heading('Payment Details', 2)
    pay_table = doc.add_table(rows=5, cols=2)
    pay_table.style = 'Light Grid Accent 1'

    pay_items = [
        ('Bank', COMPANY['bank']),
        ('Account Name', 'JW Ventures'),
        ('Account Number', '[To be provided]'),
        ('Amount', f'BND {total:.0f}'),
        ('Reference', invoice_number),
    ]

    for label, value in pay_items:
        row = pay_table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = value

    doc.add_paragraph()
    note = doc.add_paragraph('Thank you for your business!')
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].bold = True

    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph('_____________________')
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = doc.add_paragraph('Jabbar')
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.runs[0].bold = True
    title = doc.add_paragraph('Managing Director, JW Ventures Brunei')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_footer(doc)
    return doc


def create_work_order(wo_number, client_name, service_title, items, overview):
    """Create a work order with both signatures"""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_heading('WORK ORDER', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(service_title)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].bold = True

    # Info Table
    doc.add_paragraph()
    info = {
        'Work Order No:': wo_number,
        'Date:': 'March 7, 2026',
        'Client:': client_name,
        'Status:': 'PENDING APPROVAL',
    }

    table = doc.add_table(rows=2, cols=4)
    items = list(info.items())
    for i in range(4):
        if i < len(items):
            key, value = items[i]
            table.rows[0].cells[i].text = key
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = COLORS["olive"]
            table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

            table.rows[1].cells[i].text = str(value)
            table.rows[1].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Overview
    doc.add_heading('1. Scope of Work', 2)
    doc.add_paragraph(overview)

    # Services Table
    doc.add_heading('2. Services & Pricing', 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = ['Service Item', 'Description', 'Duration', 'Fee (BND)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, '8B9A6A')
        cell.paragraphs[0].runs[0].font.color.rgb = COLORS["white"]

    total = 0
    for i, item in enumerate(items):
        row_cells = table.add_row().cells
        row_cells[0].text = item.get("service", item.get("description", ""))
        row_cells[1].text = item.get("description", "")
        row_cells[2].text = item.get("duration", "")
        row_cells[3].text = item.get("fee", "")
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i % 2 == 0:
            for cell in row_cells:
                set_cell_background(cell, 'F0F0F0')

        fee_str = item.get("fee", "").replace("BND", "").strip()
        try:
            total += float(fee_str)
        except:
            pass

    # Total
    total_cells = table.add_row().cells
    total_cells[0].text = ''
    total_cells[1].text = ''
    total_cells[2].text = 'TOTAL'
    total_cells[2].paragraphs[0].runs[0].bold = True
    total_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_cells[3].text = f'BND {total:.0f}'
    total_cells[3].paragraphs[0].runs[0].bold = True
    total_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_background(total_cells[3], '8B9A6A')
    total_cells[3].paragraphs[0].runs[0].font.color.rgb = COLORS["white"]

    doc.add_paragraph()

    # Signatures (Side by Side)
    doc.add_page_break()
    doc.add_heading('3. Authorization', 1)
    doc.add_paragraph('By signing below, both parties agree to the terms and conditions outlined in this Work Order.')

    sig_table = doc.add_table(rows=8, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.columns[0].width = Inches(3.25)
    sig_table.columns[1].width = Inches(3.25)

    # Service Provider
    cell = sig_table.rows[0].cells[0]
    cell.text = 'SERVICE PROVIDER'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_background(cell, '8B9A6A')
    cell.paragraphs[0].runs[0].font.color.rgb = COLORS["white"]

    cell = sig_table.rows[1].cells[0]
    cell.text = 'JW Ventures Brunei'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)

    cell = sig_table.rows[2].cells[0]
    cell.text = 'P30000657'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].runs[0].italic = True

    sig_table.rows[3].cells[0].text = ''
    sig_table.rows[4].cells[0].text = 'Signature: _______________________'
    sig_table.rows[4].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.rows[5].cells[0].text = 'Date: ___________________'
    sig_table.rows[5].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.rows[6].cells[0].text = ''

    cell = sig_table.rows[7].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Jabbar').bold = True
    p.add_run('\nManaging Director')

    # Client
    cell = sig_table.rows[0].cells[1]
    cell.text = 'CLIENT'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_background(cell, '8B9A6A')
    cell.paragraphs[0].runs[0].font.color.rgb = COLORS["white"]

    cell = sig_table.rows[1].cells[1]
    cell.text = client_name
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)

    sig_table.rows[3].cells[1].text = ''
    sig_table.rows[4].cells[1].text = 'Signature: _______________________'
    sig_table.rows[4].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.rows[5].cells[1].text = 'Date: ___________________'
    sig_table.rows[5].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.rows[6].cells[1].text = ''

    cell = sig_table.rows[7].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('[Client Name]').bold = True
    p.add_run('\n[Title]')

    # Payment Details
    doc.add_page_break()
    doc.add_heading('4. Payment Details', 2)

    pay_table = doc.add_table(rows=5, cols=2)
    pay_table.style = 'Light Grid Accent 1'

    pay_items = [
        ('Bank', COMPANY['bank']),
        ('Account Name', 'JW Ventures'),
        ('Account Number', '[To be provided]'),
        ('Amount', f'BND {total:.0f}'),
        ('Reference', wo_number),
    ]

    for label, value in pay_items:
        row = pay_table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = value

    add_footer(doc)
    return doc


# Example usage
if __name__ == "__main__":
    # Example Invoice
    items = [
        {"description": "Domain Transfer Coordination", "qty": 1, "unit_price": "BND 100", "amount": "BND 100"},
        {"description": "DNS Configuration", "qty": 1, "unit_price": "BND 50", "amount": "BND 50"},
        {"description": "Domain Registration (1st Year)", "qty": 1, "unit_price": "BND 15", "amount": "BND 15"},
    ]

    doc = create_invoice("INV-2026-001", "Racquet & Co (PickleballBN)", "Chia Kok Yin", items)
    doc.save("test_invoice.docx")
    print("Invoice created: test_invoice.docx")
